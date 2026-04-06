from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from .models import CandidateProfile, TailoredResume
from .utils import ensure_dir, slugify


def build_resume_markdown(profile: CandidateProfile, tailored: TailoredResume) -> str:
    lines: list[str] = [
        f"# {profile.full_name}",
        profile.contact_line,
        "",
        "## Summary",
        tailored.summary,
        "",
        "## Skills",
        ", ".join(tailored.skills),
        "",
        "## Experience",
    ]
    for index in tailored.experience_indices:
        if index >= len(profile.experience):
            continue
        item = profile.experience[index]
        lines.append(f"### {item.title} | {item.company}")
        lines.append(f"{item.start} - {item.end} | {item.location}".strip(" |"))
        for bullet in item.bullets:
            lines.append(f"- {bullet}")
        lines.append("")
    if tailored.project_indices:
        lines.extend(["## Projects", ""])
        for index in tailored.project_indices:
            if index >= len(profile.projects):
                continue
            project = profile.projects[index]
            heading = project.name
            if project.technologies:
                heading = f"{heading} | {', '.join(project.technologies)}"
            lines.append(f"### {heading}")
            for bullet in project.bullets:
                lines.append(f"- {bullet}")
            if project.link:
                lines.append(f"- Link: {project.link}")
            lines.append("")
    if profile.education:
        lines.extend(["## Education", ""])
        for item in profile.education:
            lines.append(f"### {item.degree} | {item.school}")
            lines.append(f"{item.graduation} | {item.location}".strip(" |"))
            for bullet in item.bullets:
                lines.append(f"- {bullet}")
            lines.append("")
    lines.extend(
        [
            "## Job Target",
            f"{tailored.role} at {tailored.company}",
            "",
            "## Notes",
            tailored.cover_note,
        ]
    )
    return "\n".join(lines).strip() + "\n"


def write_resume_files(profile: CandidateProfile, tailored: TailoredResume, output_dir: Path) -> TailoredResume:
    dated_dir = ensure_dir(output_dir / datetime.now().strftime("%Y-%m-%d"))
    stem = f"{slugify(tailored.company)}__{slugify(tailored.role)}"
    markdown_path = dated_dir / f"{stem}.md"
    docx_path = dated_dir / f"{stem}.docx"
    markdown_path.write_text(tailored.resume_markdown, encoding="utf-8")

    document = Document()
    document.add_heading(profile.full_name, 0)
    document.add_paragraph(profile.contact_line)
    document.add_heading("Summary", level=1)
    document.add_paragraph(tailored.summary)
    document.add_heading("Skills", level=1)
    document.add_paragraph(", ".join(tailored.skills))
    document.add_heading("Experience", level=1)
    for index in tailored.experience_indices:
        if index >= len(profile.experience):
            continue
        item = profile.experience[index]
        document.add_heading(f"{item.title} | {item.company}", level=2)
        document.add_paragraph(f"{item.start} - {item.end} | {item.location}".strip(" |"))
        for bullet in item.bullets:
            document.add_paragraph(bullet, style="List Bullet")
    if tailored.project_indices:
        document.add_heading("Projects", level=1)
        for index in tailored.project_indices:
            if index >= len(profile.projects):
                continue
            project = profile.projects[index]
            heading = project.name
            if project.technologies:
                heading = f"{heading} | {', '.join(project.technologies)}"
            document.add_heading(heading, level=2)
            for bullet in project.bullets:
                document.add_paragraph(bullet, style="List Bullet")
            if project.link:
                document.add_paragraph(project.link)
    if profile.education:
        document.add_heading("Education", level=1)
        for item in profile.education:
            document.add_heading(f"{item.degree} | {item.school}", level=2)
            document.add_paragraph(f"{item.graduation} | {item.location}".strip(" |"))
            for bullet in item.bullets:
                document.add_paragraph(bullet, style="List Bullet")
    document.add_heading("Job Target", level=1)
    document.add_paragraph(f"{tailored.role} at {tailored.company}")
    document.add_heading("Notes", level=1)
    document.add_paragraph(tailored.cover_note)
    document.save(docx_path)

    tailored.output_markdown_path = markdown_path
    tailored.output_docx_path = docx_path
    return tailored
